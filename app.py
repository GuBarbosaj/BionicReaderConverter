from flask import Flask, request, jsonify, send_file
from pdf2docx import Converter
import math
from docx import Document
import math
from docx import Document
import os
import spacy
import sys
from docx import Document
from docx2pdf import convert
import io
from werkzeug.utils import secure_filename

def process_word(word, ratio):
    # Calculate the number of characters to make bold
    num_chars = math.ceil(len(word) * ratio)
    
    # Split the word into the part that should be bold and the part that shouldn't
    bold_part = word[:num_chars]
    normal_part = word[num_chars:]
    
    # Create runs for the bold part and the normal part
    bold_run = (bold_part, True)
    normal_run = (normal_part, False)
    
    # Return the runs as a list
    return [bold_run, normal_run]

def process_document(file, extension, ratio):

    if extension == 'docx':
        word_doc = Document(file)
    else:
        converted_file_in_memory = io.BytesIO()
        Converter(stream=file.read()).convert(converted_file_in_memory)
        word_doc = Document(converted_file_in_memory)



    try:
        nlp = spacy.load('en_core_web_sm')
    except OSError as e:
        print("Error loading spaCy model:", e)
        sys.exit(1)

    try:
        for paragraph in word_doc.paragraphs:
            for run in paragraph.runs:
                # Check if the run contains only text
                if len(run._r.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})) == 0:
                    # This run does not contain text, likely an image or other non-text element
                    new_run = paragraph.add_run('')
                    # Copy non-text elements (such as images) to the new run
                    for child in run._r:
                        new_run._r.append(child)


                else:
                    # Process the text as normal (negrito, etc.)
                    if run.bold:
                        continue
                    
                    # Split the run text into words
                    words = run.text.split(' ')
                    words = [' ' + word if i != 0 else word for i, word in enumerate(words)]

                    new_runs = []
                    for word in words:
                        # Use spaCy to recognize the words
                        doc = nlp(word)
                        for token in doc:
                            # Bolden a ratio of the characters in the word
                            runs = process_word(token.text, ratio)
                            new_runs.extend(runs)

                    # Clear the original run text (but preserve non-text elements)
                    run.text = ''  # This will clear the text but not remove non-text elements

                    # Add new runs with the processed text
                    for text, is_bold in new_runs:
                        new_run = paragraph.add_run(text)
                        new_run.bold = is_bold

        memory_file = io.BytesIO()
        word_doc.save(memory_file)
        memory_file.seek(0)

        return memory_file


    except Exception as e:  # Catch any other unexpected errors
        print("Unexpected error processing document:", e)
        sys.exit(1)

app = Flask(__name__)

# Lista de extensões permitidas
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    """
    Verifica se o arquivo tem uma extensão permitida.
    """
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/converter', methods=['POST'])
def upload_file():
    # Verifica se os dados foram enviados corretamente
    if 'file' not in request.files:
        return jsonify({'error': 'No file sent'}), 400

    # Recupera o arquivo e o texto
    file = request.files['file']

    # Verifica se um arquivo foi selecionado
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if not allowed_file(file.filename):
            return jsonify({"error": "Unsupported file type. Only send .pdf or .docx files."}), 400

    output = process_document(file, file.filename.rsplit('.', 1)[1].lower(), 0.5)
    
    # Retorna os dados processados
    return send_file(output, as_attachment=True, download_name="output.docx")

if __name__ == '__main__':
    app.run()
